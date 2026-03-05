import copy
import pandas as pd
import re
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def apply_default_font(doc, font_name='Aptos', font_size=9):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Apply to table cells too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_table_row_height(row, height_cm):
    """
    Set exact row height for a table row
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # cm → twips
    trHeight.set(qn('w:hRule'), 'exact')

    trPr.append(trHeight)



def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """
    Set individual cell borders. Each border is a dict:
    {"sz": 4, "val": "single", "color": "000000"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is not None:
        tcPr.remove(borders)
    borders = OxmlElement('w:tcBorders')

    for border_name, border_val in zip(["top","bottom","left","right"], [top, bottom, left, right]):
        if border_val:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), border_val.get("val","single"))
            b.set(qn('w:sz'), str(border_val.get("sz",4)))
            b.set(qn('w:color'), border_val.get("color","000000"))
            borders.append(b)
    tcPr.append(borders)


def write_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, width=None, add_para = True, font_size = 12):
    # Set cell width
    if width:
        cell.width = Cm(width)
    
    # Clear default paragraph
    p = cell.paragraphs[0]
    p.clear()
    
    # Add text
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Aptos")
    p.alignment = align
    if add_para:
        cell.add_paragraph()


def set_margin():
    doc2 = Document()

        # Default style
    normal_style = doc2.styles['Normal']
    normal_style.font.name = 'Apsto'
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = 1

        # Copy first section settings
    tgt = doc2.sections[0]

    tgt.page_width = 7772400
    tgt.page_height = 10058400
    tgt.orientation = 0
    tgt.top_margin = Cm(0.48)
    tgt.bottom_margin = Cm(0.48)
    tgt.left_margin = Cm(1.2)
    tgt.right_margin = Cm(1.2)
    return doc2



def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(cell, right={"sz": 24, "val": "single", "color": "000000"})
    sz = border thickness (8 = 1pt, 24 = 3pt)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))

    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'right', 'top', 'bottom'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["val", "sz", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell_border1(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))

    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'right', 'top', 'bottom'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["val", "sz", "color"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_row_color(row, color_hex):
    """
    row: table row object
    color_hex: string like 'D9D9D9' (no # symbol)
    """
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)



def set_logo(doc2):
    # Add table
    table = doc2.add_table(rows=1, cols=2)
    table.allow_autofit = False

    
    # Column widths
    col_widths = [6, 14.5]
    # Approx half-page width each column
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(col_widths[i])
        
        if i == 0:
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            p_left = cell.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left.paragraph_format.left_indent = Cm(0.19)

            run = p_left.add_run()
            run.add_picture(
                r"app/static/app/assets/do/image.png",
                width=Cm(3.48),
                height=Cm(3.54)
            )

        elif i == 2:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            p_right = cell.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # p_right.paragraph_format.right_indent = Cm(0.19)
            customer = []
            # Address lines
            for idx, customer_ in enumerate(customer):
                run = p_right.add_run(customer_ + '\n')
                if idx == 0:
                    run.bold = True


        elif i == 1:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            p_right = cell.add_paragraph()
            p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # -------- Line 1 --------
            # Company Name (Bold, Aptos 16)
            run1 = p_right.add_run("TRANS PACIFIC LOGISTICS PTE. LTD. ")
            run1.bold = True
            run1.font.size = Pt(16)
            run1.font.name = "Aptos"
            run1._element.rPr.rFonts.set(qn('w:eastAsia'), "Aptos")

            # UEN (Aptos 10, same line, not bold)
            run2 = p_right.add_run("(UEN 202430059Z)")
            run2.font.size = Pt(10)
            run2.font.name = "Aptos"
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), "Aptos")
            p_right.paragraph_format.line_spacing = 1.15

            # -------- Line 2 --------
            p_right.add_run("\n")
            run3 = p_right.add_run(
                "519, BALESTIER ROAD, #04-02 LE SHANTIER, SINGAPORE 329852."
            )
            run3.font.size = Pt(12)
            run3.font.name = "Aptos"
            run3._element.rPr.rFonts.set(qn('w:eastAsia'), "Aptos")
            p_right.paragraph_format.line_spacing = 1.15

            # -------- Line 3 --------
            p_right.add_run("\n")
            run4 = p_right.add_run(
                "Tel: +65 6327 1158/+65 6327 1159   Fax: +65 6225 2004"
            )
            run4.font.size = Pt(12)
            run4.font.name = "Aptos"
            run4._element.rPr.rFonts.set(qn('w:eastAsia'), "Aptos")
            p_right.paragraph_format.line_spacing = 1.15
            
    # Make right border of first column bold
    set_cell_border(
        table.rows[0].cells[0],
        right={"val": "single", "sz": 16, "color": "000000"}  # 24 = thick line
    )      
    return doc2


def set_header(doc2, s_no = str):
    table = doc2.add_table(rows=1, cols=2)
    table.allow_autofit = False
    # Column widths
    col_widths = [10.5, 9.5]
    # Approx half-page width each column
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(col_widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p_right = cell.paragraphs[0]
        if i == 0:
            p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run1 = p_right.add_run("DELIVERY NOTE")
            size = 18
        elif i == 1:
            
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run1 = p_right.add_run(s_no)
            run1.font.color.rgb = RGBColor(255, 0, 0)
            size = 26

            
        run1.bold = True
        run1.font.size = Pt(size)
        run1.font.name = "Aptos"
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), "Aptos")

    
    return doc2



def set_table(doc, lst):

    table = doc.add_table(rows=0, cols=2)
    table.allow_autofit = False

    col_widths = [10.145, 10.145]
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])

    for idx, row_data in enumerate(lst):

        row = table.add_row()
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])

        if idx == 3:
            set_row_color(row, "D9D9D9") 

        left_cell = row.cells[0]
        right_cell = row.cells[1]

        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # ----- EMPTY ROW -----
        if not row_data or idx == 10 or idx == 12:
            if not row_data:
                p = left_cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(" ")
                run.font.size = Pt(12)
            else:
                p = left_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.5  # FIX: set on paragraph
                p.paragraph_format.line_after_spacing = 1

                run = p.add_run(row_data[0])
                run.font.size = Pt(12)
            # Outer border only (no center divider)
            set_cell_border1(left_cell,
                            left={"val": "single", "sz": 8, "color": "000000"},
                            top={"val": "single", "sz": 8, "color": "000000"},
                            bottom={"val": "single", "sz": 8, "color": "000000"})

            set_cell_border1(right_cell,
                            right={"val": "single", "sz": 8, "color": "000000"},
                            top={"val": "single", "sz": 8, "color": "000000"},
                            bottom={"val": "single", "sz": 8, "color": "000000"})

            continue

       

        if len(row_data) == 1:
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5  # FIX: set on paragraph
            p.paragraph_format.line_after_spacing = 1

            run = p.add_run(row_data[0])
            run.font.size = Pt(12)

        # ----- NORMAL TWO COLUMN ROW -----
        else:
            if idx == 11:
                p_left = left_cell.paragraphs[0]
                p_left.paragraph_format.line_spacing = 1.5  # FIX here too
                run_left1 = p_left.add_run('For “Trans Pacific Logistics Pte. Ltd.”\n')
                run_left1.font.size = Pt(12)
                run_left1.bold = True

                run_left2 = p_left.add_run(row_data[0])
                run_left2.font.size = Pt(12)


                p_right = right_cell.paragraphs[0]
                p_right.paragraph_format.line_spacing = 1.5  # FIX here too
                run_right1 = p_right.add_run('For “Client Received By”\n')
                run_right1.font.size = Pt(12)
                run_right1.bold = True

                run_right2 = p_right.add_run(row_data[1])
                run_right2.font.size = Pt(12)
            else:
                p_left = left_cell.paragraphs[0]
                p_left.paragraph_format.line_spacing = 1.5  # FIX here too
                run_left = p_left.add_run(row_data[0])
                run_left.font.size = Pt(12)


                p_right = right_cell.paragraphs[0]
                p_right.paragraph_format.line_spacing = 1.5  # FIX here too
                run_right = p_right.add_run(row_data[1])
                run_right.font.size = Pt(12)
            if idx == 3:
                run_left.font.bold = True
                run_right.font.bold = True

            # ----- APPLY OUTER BORDERS -----
            set_cell_border1(left_cell,
                            left={"val": "single", "sz": 8, "color": "000000"},
                            top={"val": "single", "sz": 8, "color": "000000"},
                            bottom={"val": "single", "sz": 8, "color": "000000"},
                            right={"val": "single", "sz": 8, "color": "000000"}  # center divider
                            )

            set_cell_border1(right_cell,
                            right={"val": "single", "sz": 8, "color": "000000"},
                            top={"val": "single", "sz": 8, "color": "000000"},
                            bottom={"val": "single", "sz": 8, "color": "000000"}
                            )

    return table



def set_footer(doc):

    doc.add_paragraph()

    # -------- First Table (1 column) --------
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False   # use this instead
    table.columns[0].width = Cm(20.29)  # SET COLUMN WIDTH HERE

    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1
    # First part of text
    run1 = p.add_run("“Any claims of goods damaged / lost shall be made known to ")
    run1.font.size = Pt(12)

    # Company name (size 11 and bold)
    run2 = p.add_run("TRANS PACIFIC LOGISTICS PTE. LTD.in")
    run2.font.size = Pt(11)
    run2.bold = True

    # Remaining text
    run3 = p.add_run(" writing within 4 days from the date of the goods received. "
                     "Failure to so shall be deemed as goods received in good order "
                     "and any claim whatsoever shall not be entertained.”")
    run3.font.size = Pt(12)


    doc.add_paragraph()


    # -------- Second Table (4 columns) --------
    table1 = doc.add_table(rows=1, cols=4)
    table1.autofit = False

    col_width = 5.0725
    data = [
        'White – Customer',
        'Pink – Operation',
        'Yellow – Accounts',
        'Blue – Driver'
    ]

    for i in range(4):
        table1.columns[i].width = Cm(col_width)   # SET COLUMN WIDTH HERE

        cell = table1.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(data[i])
        run.font.size = Pt(12)

    return table


def get_table_list(data):
    lst = [
            [f"Reference No: {data['ref_no']}", ""],
            [f"PO NO: {data['po_no']}", f'DATE: {data['date']}'],
            [],
            ['COLLECTION DETAIL','DELIVERY DETAIL'],
            [f'Address:\n{data['collection_details']['address']}',
            f'Address:\n{data['delivery_details']['address']}'],
            [f'Contact Person: {data['collection_details']['contact_person']}',f'Contact Person: {data['delivery_details']['contact_person']}'],
            [f'Tel: {data['collection_details']['tel']}',f'Tel: {data['delivery_details']['tel']}'],
            [f'Collection Date: {data['collection_details']['date']}',f'Delivery Date: {data['delivery_details']['date']}'],
            [f'Collection Time: {data['collection_details']['time']}',f'Delivery Time: {data['delivery_details']['time']}'],
            [f'Total Packages:{data['total_packages']}',f'Total Weight: {data['total_weight']} KGS'],
            [f'Cargo Description: \n{data['cargo_description']}\nRef : {data['cargo_ref']}'],
            [f'Type of Truck: {data['collection_details']['truck_type']}\nTruck No: {data['collection_details']['truck_number']}\nDriver: {data['collection_details']['driver']}\nDate: {data['collection_details']['date_']}', 
             f'Name: {data['delivery_details']['name']}\nID No: {data['delivery_details']['id']}\nSignature: {data['delivery_details']['signature']}\nDate: {data['delivery_details']['date_']}'],

            [f'Remarks:\n{data['remarks']}']
            

        ]
    return lst


def get_deliver_order(data):

    doc = set_margin()

    set_logo(doc)
    set_header(doc, s_no = data['serial_no'])
    lst = get_table_list(data)
    set_table(doc, lst)
    set_footer(doc)

    return doc
