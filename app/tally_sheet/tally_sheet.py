import copy
import pandas as pd
import re
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from pathlib import Path

def get_tally_logistic(doc):

    data = {'REF NO': '', 
            'PACKING DETAILS': '',
            'CONTAINER DETAILS': '',
            'TOTAL PACKAGES': '',
            'GROSS WEIGHT': '',
            'TOTAL CBM (m³)': '',
            'TOTAL': '',
            
            }
    ref_no = doc.tables[2].rows[0].cells[2].paragraphs[0].text.strip()
    total = doc.tables[4]._cells[1].text
    pkg_details = doc.tables[3].rows[-1].cells[1].paragraphs[0].text.replace('PACKING DETAILS:\n', '').strip("\n")
    shipping = doc.tables[3].rows[-1].cells[1].paragraphs[2].text.replace('SHIPPING MARK:\n', '').strip('\n')
    for line in doc.tables[3].rows[-1].cells[1].paragraphs[1].text.split("\n"):
        if line.strip():
            key, value = line.split(":")
            key = key.strip()
            value = value.lower().replace('kgs', '').replace('cbm', '')
            
            if key == 'GROSS WEIGHT':
                num = float(value.replace(',', ''))
                value = f"{num:,.2f}"
            data[key] = value.strip()
    data["TOTAL"] = f"{float(total.lower().replace('sgd', '').replace('$', '').replace(',', '').strip()):,.2f}"
    data['PACKING DETAILS'] = pkg_details
    data['CONTAINER DETAILS'] = shipping
    data['REF NO'] = ref_no

    return data

def get_tally_data(original, final):
    df1 = pd.DataFrame(original).sort_values(by='REF NO', ascending=True).reset_index(drop=True)
    df2 = pd.DataFrame(final).sort_values(by='REF NO', ascending=True).reset_index(drop=True)
    df1.rename(columns={'TOTAL': 'ORIGINAL $', 'GROSS WEIGHT': "GROSS WEIGHT (kgs)", 'TOTAL CBM (m³)': "TOTAL CBM (m³)"}, inplace=True)
    df1['FINAL $'] = df2['TOTAL']


    # Compute totals
    total_packages = df1['TOTAL PACKAGES'].str.replace(',', '').astype(int).sum()
    total_gross_weight = df1['GROSS WEIGHT (kgs)'].str.replace(',', '').astype(float).sum()
    total_volume = df1['TOTAL CBM (m³)'].str.replace(',', '').astype(float).sum()
    total_amount_o = df1['ORIGINAL $'].str.replace(',', '').astype(float).sum()
    total_amount_f = df1['FINAL $'].str.replace(',', '').astype(float).sum()

    # Format totals with commas and 2 decimals
    total_row = {
        'REF NO': '',
        'PACKING DETAILS': '',
        'CONTAINER DETAILS': 'TOTAL',
        'TOTAL PACKAGES': str(total_packages),
        'GROSS WEIGHT (kgs)': f"{total_gross_weight:,.2f}",
        'TOTAL CBM (m³)':f"{total_volume:,.2f}",
        'ORIGINAL $': f"{total_amount_o:,.2f}",
        'FINAL $': f"{total_amount_f:,.2f}"
    }
    df = pd.concat([df1, pd.DataFrame([total_row])], ignore_index=True)
    df = df.rename(columns={
    'TOTAL PACKAGES': 'TOTAL\nPACKAGES',
    'GROSS WEIGHT (kgs)': 'GROSS\nWEIGHT (kgs)',
    'TOTAL CBM (m³)': 'TOTAL\nCBM (m³)'
})
    return df



def tally_sheet_p(writer, df1: pd.DataFrame):
    workbook = writer.book
    worksheet = writer.sheets['ORDINARY']

    wrap_center_format = workbook.add_format({
        'text_wrap': True, 'align': 'left', 'border': 1, 'valign': 'top'
    })
    text_format = workbook.add_format({
        'align': 'center', 'border': 1, 'valign': 'vcenter'
    })
    num_format = workbook.add_format({
        'num_format': '#,##0.00', 'align': 'center', 'border': 1, 'valign': 'vcenter'
    })
    header_format = workbook.add_format({
        'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1, 'text_wrap': True
    })
    bold_format = workbook.add_format({
        'bold': True, 'align': 'center', 'border': 1, 'valign': 'vcenter'
    })

    # Column widths
    worksheet.set_column('A:A', 9, text_format)
    worksheet.set_column('B:B', 49, wrap_center_format)
    worksheet.set_column('C:C', 44, text_format)
    worksheet.set_column('D:F', 14, num_format)

    # Bold last row
    last_row = len(df1)
    for col_num in range(len(df1.columns)):
        worksheet.write(last_row, col_num, df1.iloc[-1, col_num], bold_format)

    # Add table
    worksheet.add_table(0, 0, len(df1), len(df1.columns)-1, {
        'columns': [{'header': col, 'header_format': header_format} for col in df1.columns],
        'style': 'Table Style Medium 23'
    })

    return writer


def tally_sheet_f(writer, df):
    workbook = writer.book
    worksheet = writer.sheets['COMMERCIAL']

    # Reuse same formats
    wrap_center_format = workbook.add_format({'text_wrap': True, 'align': 'left', 'border': 1, 'valign': 'top'})
    text_format = workbook.add_format({'align': 'center', 'border': 1, 'valign': 'vcenter'})
    num_format = workbook.add_format({'num_format': '#,##0.00', 'align': 'center', 'border': 1, 'valign': 'vcenter'})
    header_format = workbook.add_format({'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1, 'text_wrap': True})
    bold_format = workbook.add_format({'bold': True, 'align': 'center', 'border': 1, 'valign': 'vcenter'})

    # Column widths — adjust according to your full columns
    worksheet.set_column('A:A', 9, text_format)
    worksheet.set_column('B:B', 49, wrap_center_format)
    worksheet.set_column('C:C', 44, text_format)
    worksheet.set_column('D:D', 10, num_format)
    worksheet.set_column('E:E', 12, num_format)
    worksheet.set_column('F:F', 10, num_format)
    worksheet.set_column('G:G', 11, num_format)
    worksheet.set_column('H:H', 11, num_format)

    # !!! IMPORTANT: you need df here — pass it or make it global (not recommended)
    # For now assuming it's accessible — better to pass it as argument

    last_row = len(df)  # ← this 'df' is not in scope — fix needed!
    for col_num in range(len(df.columns)):
        worksheet.write(last_row, col_num, df.iloc[-1, col_num], bold_format)

    worksheet.add_table(0, 0, len(df), len(df.columns)-1, {
        'columns': [{'header': col, 'header_format': header_format} for col in df.columns],
        'style': 'Table Style Medium 23'
    })

    return writer


# We'll use your create_tally_sheet logic, slightly adapted
def create_tally_sheet(df: pd.DataFrame, output_path: str | Path):
    excel_file = str(output_path)

    with pd.ExcelWriter(excel_file, engine='xlsxwriter') as writer:
        df1 = df[['REF NO', 'PACKING DETAILS', 'CONTAINER DETAILS',
                  'TOTAL\nPACKAGES', 'GROSS\nWEIGHT (kgs)', 'TOTAL\nCBM (m³)']]

        df1.to_excel(writer, index=False, header=False, startrow=1, sheet_name='ORDINARY')
        writer = tally_sheet_p(writer, df1)

        df.to_excel(writer, index=False, header=False, startrow=1, sheet_name='COMMERCIAL')
        writer = tally_sheet_f(writer, df)



