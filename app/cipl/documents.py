from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

from pathlib import Path
from docx import Document
BASE_DIR = Path(__file__).resolve().parent.parent
file_path = BASE_DIR / "static" / "app" / "CIPL ES-00096_.docx"

def apply_default_font(doc, font_name='Arial', font_size=9):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Apply to table cells too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_table_row_height(row, height_cm):
    """
    Set exact row height for a table row
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # cm → twips
    trHeight.set(qn('w:hRule'), 'exact')

    trPr.append(trHeight)



def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """
    Set individual cell borders. Each border is a dict:
    {"sz": 4, "val": "single", "color": "000000"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is not None:
        tcPr.remove(borders)
    borders = OxmlElement('w:tcBorders')

    for border_name, border_val in zip(["top","bottom","left","right"], [top, bottom, left, right]):
        if border_val:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), border_val.get("val","single"))
            b.set(qn('w:sz'), str(border_val.get("sz",4)))
            b.set(qn('w:color'), border_val.get("color","000000"))
            borders.append(b)
    tcPr.append(borders)


def write_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, width=None, add_para=True, header = False, indent = -0.2, right_indent = -0.2, color="black" ):
    # Set cell width
    if width:
        cell.width = Cm(width)
    if header:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Clear default paragraph
    p = cell.paragraphs[0]
    p.clear()
    
    # Remove any indent
    p.paragraph_format.left_indent = Cm(indent)      
    p.paragraph_format.right_indent = Cm(right_indent)    # left indent
    p.paragraph_format.first_line_indent = Cm(0)   # first-line indent
    
    # Add text
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    # Default black, optional red
    if color.lower() == "red":
        run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)

    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")
    p.alignment = align

    # Optional extra paragraph without indent
    if add_para:
        if text:
            if 'ROW' in text:
                run.underline = True
            if  text.startswith('MDA'):
                run.underline = True
            else:
                new_p = cell.add_paragraph()
                new_p.paragraph_format.left_indent = Cm(indent)
                new_p.paragraph_format.first_line_indent = Cm(indent)




def write_amount(cell, text, width, bold=False):
    if width:
        cell.width = Cm(width)

    p = cell.paragraphs[0]

    # Remove indent
    p.paragraph_format.left_indent = Cm(-0.1)
    p.paragraph_format.right_indent = Cm(-0)
    p.paragraph_format.first_line_indent = Cm(0)

    # Add RIGHT tab stop at end of cell width
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(width - 0.25),  # <-- adjust based on your column width
        WD_TAB_ALIGNMENT.RIGHT
    )

    # Add content
    run = p.add_run("$")
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

    p.add_run("\t")  # This moves to right tab stop

    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

# --- Function to set margins and default font ---
def set_margin():
    doc2 = Document()

    # Default style
    normal_style = doc2.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(9)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = 1


    tgt = doc2.sections[0]
 

    tgt.top_margin = Cm(0.75)
    tgt.bottom_margin = Cm(0.25)
    tgt.left_margin = Cm(1)
    tgt.right_margin = Cm(1)

    return doc2


def set_logo(doc2, customer):
    # Add table
    table = doc2.add_table(rows=1, cols=2)
    table.allow_autofit = False

    
    # Column widths
    col_widths = [10, 10.5]
    # Approx half-page width each column
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(col_widths[i])
        
        if i == 0:
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            p_left = cell.add_paragraph()
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left.paragraph_format.left_indent = Cm(0.19)

            run = p_left.add_run()
            run.add_picture(
                r'app/static/app/assets/cipl/logo.png',
                width=Cm(4.85),
                height=Cm(1.30)
            )

        elif i == 1:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            p_right = cell.add_paragraph()
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # p_right.paragraph_format.right_indent = Cm(0.19)

            # Address lines
            for idx, customer_ in enumerate(customer):
                if customer_.endswith('Singapore Pte. Ltd.'):
                    customer_ = customer_.replace('Singapore Pte. Ltd.', 'Singapore Pte. Ltd')

                run = p_right.add_run(customer_ + '\n')
                if idx == 0:
                    run.bold = True
            
            
    return doc2


def set_header(doc2, text = "COMMERCIAL INVOICE / PACKING LIST"):
    table = doc2.add_table(rows=1, cols=1)
    table.allow_autofit = False
    table.autofit = False
    col_width = 21
    # Approx half-page width each column
    table.columns[0].width = Cm(col_width)
    # table.columns[1].width = Cm(9)
    row = table.rows[0]
    set_table_row_height(row, height_cm=0.60)

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(col_width)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is not None:
                tcPr.remove(tcBorders)

    # Left cell: logo
    left_cell = table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # p_left = left_cell.add_paragraph(text)
    # p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left = left_cell.paragraphs[0]   # use default empty paragraph
    p_left.text = text
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p_left.runs
    for run in p_left.runs:
        run.font.name = 'Arial'  
        run.font.size = Pt(12) 
        run.font.color.rgb = RGBColor(23, 54, 93)
        run.font.bold = True
    run = p_left.add_run()
    return doc2

def set_consignee(doc2, consginee, labels):
    # Add table
    table = doc2.add_table(rows=1, cols=3)
    table.allow_autofit = False


    def add_run(p, text, bold=False, align=None):
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")
    # Column widths
    col_widths = [10, 5.2, 5.3]
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])
    table.columns[2].width = Cm(col_widths[2])

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(col_widths[i])
        
        if i == 0:
            p = cell.paragraphs[0]
            for idx, con_ in enumerate(consginee):
                bold = False
                if con_.endswith('189/388 Village No. 5, Pharam 2Rd, Panthai Norasing, Muang,'):
                    con_ = con_.replace('189/388 Village No. 5, Pharam 2Rd, Panthai Norasing, Muang,', '189/388 Village No. 5, Pharam 2Rd, Panthai')
                if con_.endswith('Samut Sakhon 74000'):
                    con_ = con_.replace('Samut Sakhon 74000', 'Norasing, Muang, Samut Sakhon 74000')
                if idx == 0:
                    bold = True
                add_run(p, con_ +"\n", bold=bold)
            
           
        elif i == 1:
            for i in range(7):
                p_mid = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                add_run(p_mid, list(labels.keys())[i], bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

        else:
            for i in range(7):
                p_right = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                add_run(p_right, list(labels.values())[i], align=WD_ALIGN_PARAGRAPH.RIGHT)
                # p_right.paragraph_format.right_indent = Cm(0.19)
   
    
    # doc2.add_paragraph()
    return doc2

def add_packing_details(doc, col_widths, packing_details, indent):

    table = doc.tables[-1]

    cells = table.add_row().cells
    for i, cell in enumerate(cells):
        write_cell(cells[i], "", align=WD_ALIGN_PARAGRAPH.LEFT, width=col_widths[i], add_para= False)
        set_cell_borders(
            cell,
            left={"val":"single","sz":8,"color":"000000"},
            right={"val":"single","sz":8,"color":"000000"},
            top=None,
            bottom=None
        )
    col_idx = 1
    total_table_lines = 48
    add_next_line = packing_details['present_lines']
    if add_next_line >0:
        packing_header = f"{"\n" *( total_table_lines - add_next_line)}PACKING DETAILS:\n"
    else:
        packing_header = "PACKING DETAILS:\n"
       
    p = cells[1].paragraphs[0]

    # Remove any indent
    p.paragraph_format.left_indent = Cm(0)   # no left indent
    p.paragraph_format.right_indent = Cm(-0.2)  # no right indent
    p.paragraph_format.first_line_indent = Cm(0)  # no first-line indent

    # Add header
    run = p.add_run(packing_header)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

    # Add each detail
    for d in packing_details["details"]:
        run = p.add_run(d + "\n")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")   

    pkg = cells[1].add_paragraph()

    pkg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pkg.paragraph_format.right_indent = Cm(-0.2)

        
    for idx, d in enumerate(packing_details['total'].items()):
        key, val = d
        pkg.add_run(f"{key} {val}\n").bold = True


    for run in pkg.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

    pkg1 = cells[1].add_paragraph()

    pkg1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pkg1.paragraph_format.right_indent = Cm(0.19)
    pkg1.add_run("SHIPPING MARK:\n").bold = True
    for data in packing_details['shipping']:
        pkg1.add_run(f"{data}\n")



    for run in pkg1.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")


    return doc

def add_total_row(doc, col_width, total):
    
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.allow_autofit = False
    footer_table.autofit = False

    # Set table left indent to 0.19 cm
    tbl_pr = footer_table._tbl.tblPr
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), str(Cm(0.19).twips))
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)

    col_widths = [col_width.pop(), sum(col_width) ]

    
    footer_table.columns[0].width = Cm(col_widths[1])
    footer_table.columns[1].width = Cm(col_widths[0])


    # # Apply solid border to both cells
    for idx, cell in enumerate(footer_table.rows[0].cells):
        if idx == 1:
            number = float(total.replace(',', ''))
            total = f"{number:,.2f}"
            write_amount(cell, total, width=col_widths[0], bold= True)
        if idx == 0:
            write_cell(cell, "TOTAL INVOICE VALUE:", align=WD_ALIGN_PARAGRAPH.RIGHT, width=col_widths[1], bold= True, add_para = False, right_indent= 0)

        set_cell_borders(
            cell,
            top={"val":"single","sz":8,"color":"000000"},
            bottom={"val":"single","sz":8,"color":"000000"},
            left={"val":"single","sz":8,"color":"000000"},
            right={"val":"single","sz":8,"color":"000000"},
        )
    tbl = footer_table._tbl
    tblPr = tbl.tblPr

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    return doc



def add_items_table(doc, items, packing_details, total):
    headers = [
        "ITEM NO.", "DESCRIPTION", "SERIAL NO.", "UOM", "COO", "QTY", "UNIT COST", "TOTAL - SGD $"
    ]


    # Add table
    table = doc.add_table(rows=1, cols=len(headers))
    table.allow_autofit = False

    # Set table left indent to 0.19 cm
    tbl_pr = table._tbl.tblPr
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), str(Cm(0.19).twips))
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)

    # Column widths
    
    col_widths = [1.11, 9.17,1.55,1.11,1.11, 1.11, 1.99,2.43]
    # Header row
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, width=col_widths[i], add_para= False, header = True)

    # Data rows
    for row in items:
        cells = table.add_row().cells
        for i, val in enumerate(row.values()):
            if i  >7:
                break
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            indent = 0 if i == 1 else -0.2
            if i >= 6 and val:
                write_amount(cells[i], val.replace('$','').strip(), width=col_widths[i])
                continue
            if i == 1 and val and row['description_modified'] != 1:

                write_cell(cells[i], val, align=align, width=col_widths[i], indent= indent, color = 'red')
            else:
                write_cell(cells[i], val, align=align, width=col_widths[i], indent= indent)
            
    num_rows = len(items)
    for idx, row in enumerate(table.rows):
        if idx == 0:
            top={"val":"single","sz":8,"color":"000000"}
            bottom={"val":"single","sz":8,"color":"000000"}
        elif idx == num_rows:
            top=None
            bottom=None
        else:
            top=None
            bottom=None
        for cell in row.cells:
            set_cell_borders(
                cell,
                left={"val":"single","sz":8,"color":"000000"},
                right={"val":"single","sz":8,"color":"000000"},
                top=top,
                bottom=bottom
            )
    
    # Optional: set header row height
    table.rows[0].height = Cm(0.8)
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    add_packing_details(doc, col_widths, packing_details, indent)

    add_total_row(doc, col_widths, total)
    doc.add_paragraph()

    return doc


def set_footer(doc2):
    # Add table
    table = doc2.add_table(rows=1, cols=1)
    table.allow_autofit = False

    for i, cell in enumerate(table.rows[0].cells):
        
        if i == 0:
            
            pkg = cell.paragraphs[0]
            pkg.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pkg.paragraph_format.right_indent = Cm(0.19)
            pkg.add_run("For and on behalf of: (Exentec Malaysia Pte Ltd)\n\n\n").bold = True
            pkg.add_run("________________________________________________\n")
            pkg.add_run("Authorised Signature")

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("Page 1")

    for run in pkg.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")


    return doc2


def create_documents(data, filename):
    # ───────────────────────────────────────────────
    # 1. Load original & create new doc (your code)
    # ───────────────────────────────────────────────

    doc = set_margin()
    # ───────────────────────────────────────────────
    # 2. Add invisible table for left/right layout
    # ───────────────────────────────────────────────
    doc = set_logo(doc, customer= data['shipper'])
    # apply_default_font(doc, font_name='Arial', font_size=9)
    doc = set_header(doc)
    doc = set_consignee(doc, consginee= data['importer_of_record'], labels= data['labels'])
    doc = add_items_table(doc, items= data['items'], packing_details = data['packing_details'], total = data['total'])

    doc = set_footer(doc)
    return doc
    doc.save("output.docx")


