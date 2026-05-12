import copy
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from math import floor
import os


def get_overall_total(ori, fnl):
    data = {'ori':[], 'fnl':[]}
    perc_10 = floor((float(ori.iloc[-1]['10%'].replace(',', ''))))
    perc_7 = floor((float(ori.iloc[-1]['7%'].replace(',', ''))))
    total = perc_10 + perc_7
    data['ori'].extend([f'{perc_10:,.2f}', f'{perc_7:,.2f}', f'{total:,.2f}'])

    perc_10 = floor((float(fnl.iloc[-1]['10%'].replace(',', ''))))
    perc_7 = floor((float(fnl.iloc[-1]['7%'].replace(',', ''))))
    total = perc_10 + perc_7
    data['fnl'].extend([f'{perc_10:,.2f}', f'{perc_7:,.2f}', f'{total:,.2f}'])
    return data


# alignment para
def update_page_one(table, ori, fnl, convert = 'ORI', verification = True):
    data = get_overall_total(ori, fnl)

    if convert == 'ORI':
        ori_values = data['ori']
        fnl_values = data['fnl']
    else:

        ori_values = data['fnl']
        fnl_values = data['ori']
    row_idx, cell_idx, para_idx = 0, 0, 0
    for idx, old_value in enumerate(fnl_values):
        new_value = ori_values[idx]
        
        for r_idx, row in enumerate(table.rows[row_idx:]):
            for c_idx, cell in enumerate(row.cells[cell_idx:]):
                for p_idx, paragraph in enumerate(cell.paragraphs[para_idx:]):
                        
                        if old_value in paragraph.text:
                           
    #                      # Rewrite using your function
                            write_amount(
                                cell,
                                p_idx,
                                width = 3.2,
                                text = new_value,
                                bold = True if idx ==2 else False,
                                verification= verification
                            )
                            row_idx += r_idx
                            cell_idx += c_idx
                            para_idx += para_idx
                            break
    # update_page_thb(table, ori, fnl, convert, row_idx, cell_idx, para_idx)
    

# alignment para
def update_page_thb(table, ori, fnl, convert = 'ORI', row_idx = 0, cell_idx = 0, para_idx = 0, verification = True):
  
    if convert == 'ORI':
        original = ori
        final = fnl
    else:
        original = fnl
        final = ori

    columns = ['AMOUNT', 'THB', '10%', 'THB+10%', '7%', 'PAGE TOTAL']
    # table_text = [paragraph.text for row in doc.tables[0].rows for cell in row.cells for paragraph in cell.paragraphs]
    for col in columns:
        for idx in range(original.shape[0]):
            new_value, old_value = original[col][idx], final[col][idx]
            if pd.notna(old_value) and pd.notna(new_value):
                # print(col, old_value, new_value)
                width = 2.5 if col in ['THB+10%', '7%'] else 2.85 if col in 'THB' else 2.6
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, paragraph in enumerate(cell.paragraphs):
                # for row in table.rows:
                #    for cell in row.cells:
                #     for p_idx, paragraph in enumerate(cell.paragraphs):
                            if old_value in paragraph.text:
                                print(col, old_value, new_value)
                                if col == 'AMOUNT':
                                    # remaining_words = table_text.count(f'SGD\t{old_value}') -1
                                    # Rewrite using your function
                                    write_amount(
                                        cell,
                                        p_idx,
                                        width = 2.85,
                                        text = new_value,
                                        verification= verification,
                                        sgd=True   # or True if needed
                                    )
                                elif col == 'PAGE TOTAL':
                                    # remaining_words = table_text.count(old_value) -1
                                    write_amount(
                                        cell,
                                        p_idx,
                                        text = new_value,
                                        width = 4.5,
                                        fontsize=11,
                                        bold=True,
                                        verification= verification
                                    )
                                else:
                                    # remaining_words = table_text.count(old_value) -1
                                    write_amount(
                                        cell,
                                        p_idx,
                                        width = width,
                                        text = new_value,
                                        verification= verification,
                                        # or True if needed
                                    )
                                # if remaining_words ==0:
                                #     print(f'{col}--> {old_value} --> {new_value} --> break')
                                #     row_idx += r_idx
                                #     cell_idx += c_idx
                                #     para_idx += para_idx
                                #     break

                              

def write_amount(cell, p_idx, text, fontsize = 8, width = 2.5, bold=False, sgd = False, verification = True):

    if width:
        cell.width = Cm(width)

    p = cell.paragraphs[p_idx]
    p.clear()

    # Remove indent
    p.paragraph_format.left_indent = Cm(0.1)
    p.paragraph_format.right_indent = Cm(-0)
    p.paragraph_format.first_line_indent = Cm(0)

    # Add RIGHT tab stop at end of cell width
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(width),  # <-- adjust based on your column width
        WD_TAB_ALIGNMENT.RIGHT
    )

    if sgd:
        # Add content
        run = p.add_run("SGD")
        run.bold = bold
        run.font.name = "Arial"
        if verification:
            run.font.color.rgb = RGBColor(0, 128, 0)
        run.font.size = Pt(fontsize)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

    p.add_run("\t")  # This moves to right tab stop

    run = p.add_run(f"{text}")
    run.bold = bold
    run.font.name = "Arial"
    if verification:
        run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(fontsize)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")


async def create_thai_document(ori, fnl, doc_path, index_, verification, fnl_to_ori):

    doc = Document(doc_path)
    convert = 'ORI' if fnl_to_ori else 'FNL'
    for idx, table in enumerate(doc.tables):
        if idx ==0:
            update_page_one(table, ori, fnl, convert = 'ORI', verification = verification)
            update_page_thb(table, ori[0:index_[0]+1], fnl[0:index_[0]+1], convert, verification = verification)
        else:
            update_page_thb(table, ori[index_[idx-1] +1:index_[idx]+1].reset_index(drop=True), fnl[index_[idx-1] +1:index_[idx]+1].reset_index(drop=True), convert, verification = verification)

    return doc
