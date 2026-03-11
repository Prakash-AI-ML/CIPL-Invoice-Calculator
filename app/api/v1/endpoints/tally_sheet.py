from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List
import os
import tempfile
import shutil
from pathlib import Path
from docx import Document
import pandas as pd
from fastapi.security import HTTPAuthorizationCredentials
from fastapi.responses import RedirectResponse, JSONResponse
from sqlalchemy.ext.asyncio import AsyncSession

import uuid
from typing import List
import io
import pandas as pd
from fastapi.responses import StreamingResponse
from fastapi import FastAPI, UploadFile, Request, HTTPException, Header
from typing import List, Optional
import json
import logging
from datetime import datetime
from app.cipl.extract import analysis_cipl
from app.cipl.documents import create_documents
from app.db.session import get_db
from app.api.dependencies import get_user_permissions_detailed
from app.crud.log_manage import backend_logs
from app.tally_sheet.tally_sheet import *
import io, uuid, zipfile, logging
from docx2pdf import convert
import traceback
import tempfile

# ─── FastAPI app ────────────────────────────────────────

router = APIRouter() 


@router.post("/generate1")
async def generate_tally_sheet(
    original_files: List[UploadFile] = File(...),
    final_files: List[UploadFile] = File(...)
):
    """
    Upload ORIGINAL and FINAL .docx files → generates TALLY_SHEET.xlsx
    """
    if not original_files or not final_files:
        raise HTTPException(400, "Both original_files and final_files are required")

    # Create temporary directory
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)

        orig_dir = tmp_path / "ORIGINAL" / "WORD"
        final_dir = tmp_path / "FINAL" / "WORD"

        orig_dir.mkdir(parents=True, exist_ok=True)
        final_dir.mkdir(parents=True, exist_ok=True)

        # Save uploaded files
        for file in original_files:
            if not file.filename.lower().endswith('.docx'):
                raise HTTPException(400, f"Only .docx files allowed. Got: {file.filename}")
            dest = orig_dir / file.filename
            with dest.open("wb") as f:
                shutil.copyfileobj(file.file, f)

        for file in final_files:
            if not file.filename.lower().endswith('.docx'):
                raise HTTPException(400, f"Only .docx files allowed. Got: {file.filename}")
            dest = final_dir / file.filename
            with dest.open("wb") as f:
                shutil.copyfileobj(file.file, f)

        # ─── Your original logic, adapted ───────────────────────
        original = []
        for file in orig_dir.iterdir():
            if file.is_file() and file.suffix.lower() == '.docx':
                doc = Document(file)
                original.append(get_tally_logistic(doc))

        final_list = []  # renamed to avoid shadowing built-in 'final'
        for file in final_dir.iterdir():
            if file.is_file() and file.suffix.lower() == '.docx':
                doc = Document(file)
                final_list.append(get_tally_logistic(doc))

        df = get_tally_data(original, final_list)
        print(df)

        # Output file
        output_excel = tmp_path / "TALLY_SHEET.xlsx"
        create_tally_sheet(df, output_excel)

        # Return the file
        return FileResponse(
            path=output_excel,
            filename="TALLY_SHEET.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


@router.post("/generate")
async def generate_tally_sheet(
    original_files: List[UploadFile] = File(...),
    final_files: List[UploadFile] = File(...)
):
    if not original_files or not final_files:
        raise HTTPException(400, "Both original_files and final_files are required")

    # We'll keep the temp directory until AFTER we send the response
    # → use manual cleanup instead of with-statement
    tmp_dir = None
    try:
        tmp_dir = Path(tempfile.mkdtemp(prefix="tally_"))
        print(f"Created temp dir: {tmp_dir}")

        orig_dir = tmp_dir / "ORIGINAL" / "WORD"
        final_dir = tmp_dir / "FINAL" / "WORD"
        orig_dir.mkdir(parents=True, exist_ok=True)
        final_dir.mkdir(parents=True, exist_ok=True)

        # Save uploaded files
        for file in original_files:
            if not file.filename.lower().endswith(('.docx', '.doc')):
                raise HTTPException(400, f"Only .docx files allowed. Got: {file.filename}")
            dest = orig_dir / file.filename
            with dest.open("wb") as f:
                shutil.copyfileobj(file.file, f)

        for file in final_files:
            if not file.filename.lower().endswith(('.docx', '.doc')):
                raise HTTPException(400, f"Only .docx files allowed. Got: {file.filename}")
            dest = final_dir / file.filename
            with dest.open("wb") as f:
                shutil.copyfileobj(file.file, f)

        # ─── Process documents ────────────────────────────────
        original_data = []
        for file in orig_dir.glob("*.docx"):
            try:
                doc = Document(file)
                original_data.append(get_tally_logistic(doc))
            except Exception as e:
                print(f"Error reading original file {file.name}: {e}")

        final_data = []
        for file in final_dir.glob("*.docx"):
            try:
                doc = Document(file)
                final_data.append(get_tally_logistic(doc))
            except Exception as e:
                print(f"Error reading final file {file.name}: {e}")

        if not original_data or not final_data:
            raise HTTPException(400, "No valid data extracted from uploaded Word files")

        df = get_tally_data(original_data, final_data)
        print("DataFrame shape:", df.shape)
        # print(df.head().to_string())

        # Output file
        output_excel = os.path.join(tmp_dir, "TALLY_SHEET.xlsx")

        # This is the critical part — catch & log any error
        try:
            create_tally_sheet(df, output_excel)
        except Exception as exc:
            print("ERROR during create_tally_sheet:")
            traceback.print_exc()
            raise HTTPException(
                status_code=500,
                detail=f"Failed to generate Excel file: {str(exc)}"
            )

        # if not output_excel.is_file():
        #     raise HTTPException(
        #         500,
        #         "Excel file was not created — check logs for errors in create_tally_sheet"
        #     )

        # Send file — file will be kept open until response is sent
        return FileResponse(
            path=output_excel,
            filename="TALLY_SHEET.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        traceback.print_exc()

    except Exception as e:
        raise HTTPException(500, f"Server error: {str(e)}")

    # finally:
    #     # Clean up temp folder — runs even if exception occurred
    #     if tmp_dir is not None and tmp_dir.exists():
    #         try:
    #             shutil.rmtree(tmp_dir, ignore_errors=True)
    #             print(f"Cleaned up: {tmp_dir}")
    #         except Exception as cleanup_err:
    #             print(f"Cleanup failed: {cleanup_err}")

